#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
数据导入导出模块
支持TXT、Excel、Word、PDF等多种格式
"""

import os
import json
from typing import List, Dict, Any, Optional

# 尝试导入必要的库

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

from .utils import get_file_extension, ensure_dir_exists, check_file_exists


class DataIO:
    """数据导入导出类"""
    
    @staticmethod
    def import_from_txt(file_path: str, encoding: str = 'utf-8', delimiter: str = None) -> List[Dict[str, Any]]:
        """
        从TXT文件导入数据
        
        Args:
            file_path: 文件路径
            encoding: 文件编码
            delimiter: 分隔符，None表示按行读取
        
        Returns:
            List[Dict[str, Any]]: 导入的数据列表
        """
        if not check_file_exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        data = []
        with open(file_path, 'r', encoding=encoding) as f:
            lines = f.readlines()
        
        if delimiter:
            # 按分隔符读取，第一行为表头
            if len(lines) > 0:
                headers = [h.strip() for h in lines[0].split(delimiter)]
                for line in lines[1:]:
                    if line.strip():
                        values = [v.strip() for v in line.split(delimiter)]
                        if len(values) == len(headers):
                            data.append(dict(zip(headers, values)))
        else:
            # 按行读取
            for i, line in enumerate(lines):
                line = line.strip()
                if line:
                    data.append({"content": line, "index": i+1})
        
        return data
    
    @staticmethod
    def import_from_excel(file_path: str, sheet_name: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        从Excel文件导入数据
        
        Args:
            file_path: 文件路径
            sheet_name: 工作表名称，None表示第一个工作表
        
        Returns:
            List[Dict[str, Any]]: 导入的数据列表
        """
        if pd is None:
            raise ImportError("请安装pandas库: pip install pandas openpyxl")
        
        if not check_file_exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        df = pd.read_excel(file_path, sheet_name=sheet_name)
        data = df.to_dict('records')
        
        return data
    
    @staticmethod
    def import_data(file_path: str, **kwargs) -> List[Dict[str, Any]]:
        """
        自动根据文件扩展名导入数据
        
        Args:
            file_path: 文件路径
            **kwargs: 额外参数
        
        Returns:
            List[Dict[str, Any]]: 导入的数据列表
        """
        ext = get_file_extension(file_path)
        
        if ext == '.txt':
            return DataIO.import_from_txt(file_path, **kwargs)
        elif ext in ['.xlsx', '.xls']:
            return DataIO.import_from_excel(file_path, **kwargs)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    @staticmethod
    def export_to_txt(data: List[Dict[str, Any]], file_path: str, encoding: str = 'utf-8', delimiter: str = None) -> None:
        """
        导出数据到TXT文件
        
        Args:
            data: 要导出的数据列表
            file_path: 文件路径
            encoding: 文件编码
            delimiter: 分隔符，None表示按行导出
        """
        ensure_dir_exists(os.path.dirname(file_path))
        
        with open(file_path, 'w', encoding=encoding) as f:
            if delimiter and data:
                # 带表头的分隔符格式
                headers = list(data[0].keys())
                f.write(delimiter.join(headers) + '\n')
                
                for item in data:
                    values = [str(item.get(header, '')) for header in headers]
                    f.write(delimiter.join(values) + '\n')
            else:
                # 按行导出
                for item in data:
                    if isinstance(item, dict):
                        f.write(str(item.get('content', str(item))) + '\n')
                    else:
                        f.write(str(item) + '\n')
    
    @staticmethod
    def export_to_excel(data: List[Dict[str, Any]], file_path: str, sheet_name: str = 'Sheet1') -> None:
        """
        导出数据到Excel文件
        
        Args:
            data: 要导出的数据列表
            file_path: 文件路径
            sheet_name: 工作表名称
        """
        if pd is None:
            raise ImportError("请安装pandas库: pip install pandas openpyxl")
        
        ensure_dir_exists(os.path.dirname(file_path))
        
        df = pd.DataFrame(data)
        df.to_excel(file_path, sheet_name=sheet_name, index=False)
    
    @staticmethod
    def export_to_word(data: List[Dict[str, Any]], file_path: str, title: str = '') -> None:
        """
        导出数据到Word文件
        
        Args:
            data: 要导出的数据列表
            file_path: 文件路径
            title: 文档标题
        """
        if Document is None:
            raise ImportError("请安装python-docx库: pip install python-docx")
        
        ensure_dir_exists(os.path.dirname(file_path))
        
        doc = Document()
        
        # 添加标题
        if title:
            doc.add_heading(title, level=1)
            doc.add_paragraph()
        
        # 添加数据
        for i, item in enumerate(data, 1):
            if isinstance(item, dict):
                # 对于字典类型，按键值对添加
                doc.add_heading(f"项目 {i}", level=2)
                for key, value in item.items():
                    doc.add_paragraph(f"{key}: {value}")
            else:
                # 对于其他类型，直接添加
                doc.add_paragraph(str(item))
            
            doc.add_paragraph()  # 添加空行分隔
        
        doc.save(file_path)
    
    @staticmethod
    def export_to_pdf(data: List[Dict[str, Any]], file_path: str, title: str = '') -> None:
        """
        导出数据到PDF文件
        
        Args:
            data: 要导出的数据列表
            file_path: 文件路径
            title: 文档标题
        """
        if FPDF is None:
            raise ImportError("请安装fpdf库: pip install fpdf")
        
        ensure_dir_exists(os.path.dirname(file_path))
        
        pdf = FPDF()
        pdf.add_page()
        
        # 设置字体
        pdf.set_font("Arial", size=12)
        
        # 添加标题
        if title:
            pdf.set_font("Arial", 'B', size=16)
            pdf.cell(200, 10, txt=title, ln=1, align='C')
            pdf.ln(10)
            pdf.set_font("Arial", size=12)
        
        # 添加数据
        for i, item in enumerate(data, 1):
            if isinstance(item, dict):
                # 对于字典类型，按键值对添加
                pdf.set_font("Arial", 'B', size=14)
                pdf.cell(200, 10, txt=f"项目 {i}", ln=1)
                pdf.set_font("Arial", size=12)
                
                for key, value in item.items():
                    pdf.cell(50, 10, txt=f"{key}: ", ln=0)
                    pdf.multi_cell(150, 10, txt=str(value))
            else:
                # 对于其他类型，直接添加
                pdf.multi_cell(0, 10, txt=str(item))
            
            pdf.ln(5)  # 添加空行分隔
        
        pdf.output(file_path)
    
    @staticmethod
    def export_data(data: List[Dict[str, Any]], file_path: str, **kwargs) -> None:
        """
        自动根据文件扩展名导出数据
        
        Args:
            data: 要导出的数据列表
            file_path: 文件路径
            **kwargs: 额外参数
        """
        ext = get_file_extension(file_path)
        
        if ext == '.txt':
            DataIO.export_to_txt(data, file_path, **kwargs)
        elif ext in ['.xlsx', '.xls']:
            DataIO.export_to_excel(data, file_path, **kwargs)
        elif ext == '.docx':
            DataIO.export_to_word(data, file_path, **kwargs)
        elif ext == '.pdf':
            DataIO.export_to_pdf(data, file_path, **kwargs)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    @staticmethod
    def save_to_json(data: Any, file_path: str, encoding: str = 'utf-8') -> None:
        """
        保存数据到JSON文件
        
        Args:
            data: 要保存的数据
            file_path: 文件路径
            encoding: 文件编码
        """
        ensure_dir_exists(os.path.dirname(file_path))
        
        with open(file_path, 'w', encoding=encoding) as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    
    @staticmethod
    def load_from_json(file_path: str, encoding: str = 'utf-8') -> Any:
        """
        从JSON文件加载数据
        
        Args:
            file_path: 文件路径
            encoding: 文件编码
        
        Returns:
            Any: 加载的数据
        """
        if not check_file_exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        with open(file_path, 'r', encoding=encoding) as f:
            return json.load(f)
    
    @staticmethod
    def batch_import(file_paths: List[str], **kwargs) -> List[Dict[str, Any]]:
        """
        批量导入数据
        
        Args:
            file_paths: 文件路径列表
            **kwargs: 额外参数
        
        Returns:
            List[Dict[str, Any]]: 导入的数据列表
        """
        all_data = []
        
        for file_path in file_paths:
            data = DataIO.import_data(file_path, **kwargs)
            all_data.extend(data)
        
        return all_data
    
    @staticmethod
    def batch_export(data: List[Dict[str, Any]], output_dir: str, file_format: str = 'txt', **kwargs) -> List[str]:
        """
        批量导出数据
        
        Args:
            data: 要导出的数据列表
            output_dir: 输出目录
            file_format: 文件格式
            **kwargs: 额外参数
        
        Returns:
            List[str]: 导出的文件路径列表
        """
        ensure_dir_exists(output_dir)
        
        file_paths = []
        for i, item in enumerate(data, 1):
            file_path = os.path.join(output_dir, f"item_{i}.{file_format}")
            DataIO.export_data([item], file_path, **kwargs)
            file_paths.append(file_path)
        
        return file_paths
    
    @staticmethod
    def validate_data(data: List[Dict[str, Any]], required_fields: List[str]) -> bool:
        """
        验证数据是否包含所有必填字段
        
        Args:
            data: 要验证的数据列表
            required_fields: 必填字段列表
        
        Returns:
            bool: 数据是否有效
        """
        for item in data:
            if not all(field in item for field in required_fields):
                return False
        
        return True
